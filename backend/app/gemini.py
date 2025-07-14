from fastapi import APIRouter, BackgroundTasks, UploadFile, File, Form, HTTPException
from uuid import uuid4
import os
import re
import time
from pathlib import Path
from dotenv import load_dotenv
from supabase import create_client, Client
import google.generativeai as genai
import pdfplumber
import json
from docx import Document
import io
import requests
from datetime import datetime
from PIL import Image
import pytesseract

router = APIRouter()

# ENV VARS
env_path = Path(__file__).parent / ".env"
load_dotenv(dotenv_path=env_path)
SUPABASE_URL = os.getenv("SUPABASE_URL")
SUPABASE_SERVICE_ROLE_KEY = os.getenv("SUPABASE_SERVICE_ROLE_KEY")
supabase: Client = create_client(SUPABASE_URL, SUPABASE_SERVICE_ROLE_KEY)
genai.configure(api_key=os.getenv("GOOGLE_API_KEY"))
pytesseract.pytesseract.tesseract_cmd = r"C:\Tesseract\tesseract.exe"
os.environ['TESSDATA_PREFIX'] = r"C:\Tesseract"


@router.post("/upload-cv")
async def upload_cv(user_id: str = Form(...), file: UploadFile = File(...)):
    ext = os.path.splitext(file.filename)[-1].lower()
    filename = f"{user_id}/cv_{int(time.time())}_{file.filename}"

    result = supabase.storage.from_("user-uploads").upload(filename, await file.read(), upsert=True)

    if result.error:
        raise HTTPException(status_code=500, detail="Greška pri uploadu.")

    signed = supabase.storage.from_("user-uploads").create_signed_url(filename, 3600 * 24)
    if signed.error:
        raise HTTPException(status_code=500, detail="Greška pri kreiranju signed URL-a")

    supabase.table("users").update({"cv_url": signed.data["signedUrl"]}).eq("id", user_id).execute()

    try:
        analysis = find_my_jobs(user_id)
        supabase.table("users").update({
            "job_keywords": analysis["keywords"],
            "job_category": analysis["category"],
            "job_analysis_last_updated": datetime.utcnow().isoformat()
        }).eq("id", user_id).execute()
    except Exception as e:
        print("⚠️ Greška u automatskoj analizi:", str(e))

    return {"success": True, "cv_url": signed.data["signedUrl"]}


@router.get("/user-job-analysis/{user_id}")
def get_user_job_analysis(user_id: str):
    try:
        print("🔍 FETCHING ANALYSIS for user:", user_id)

        user = supabase.table("users").select("job_category, job_keywords").eq("id", user_id).single().execute()
        print("👤 USER RAW:", user)

        if not user.data:
            print("❌ Nema korisničkih podataka u bazi.")
            raise HTTPException(status_code=404, detail="Nema podataka")

        category = user.data.get("job_category")
        keywords = user.data.get("job_keywords")

        print("✅ category:", category)
        print("✅ keywords:", keywords)

        if not category or not keywords:
            raise HTTPException(status_code=404, detail="Nema spremljene analize")

        jobs_response = supabase.table("jobs").select("*").ilike("job_type", category).execute()
        jobs = jobs_response.data if jobs_response.data else []

        def match_score(job):
            combined = f"{job['title']} {job['description']}".lower()
            return sum(1 for kw in keywords if kw.lower() in combined)

        ranked = sorted(jobs, key=match_score, reverse=True)

        return {
            "category": category,
            "keywords": keywords,
            "results": ranked[:10]
        }

    except Exception as e:
        print("🔥 BACKEND ERROR:", str(e))
        raise HTTPException(status_code=500, detail=f"Greška: {str(e)}")


@router.post("/analyze-cv/{user_id}/{job_id}")
def analyze_cv(user_id: str, job_id: str, background_tasks: BackgroundTasks):
    analysis_id = str(uuid4())
    
    user_lookup = supabase.table("users").select("id").eq("id", user_id).single().execute()
    if not user_lookup.data:
        raise HTTPException(status_code=404, detail="Korisnik nije pronađen")

    real_user_id = user_lookup.data["id"]

    supabase.table("application_analysis").insert({
        "id": analysis_id,
        "user_id": real_user_id,
        "job_id": job_id,
        "analysis": None,
        "score": None
    }).execute()

    background_tasks.add_task(run_analysis_task, analysis_id)
    return {"analysis_id": analysis_id}




def run_analysis_task(analysis_id: str):
    result = supabase.table("application_analysis").select("*").eq("id", analysis_id).execute()
    record = result.data[0] if result.data else None
    if not record:
        return

    user_id = record["user_id"]
    job_id = record["job_id"]

    for _ in range(5):
        job_data = supabase.from_("jobs").select("description").eq("id", job_id).single().execute()
        if job_data.data and job_data.data.get("description"):
            break
        time.sleep(1)

    user_data = supabase.table("users").select("cv_url").eq("id", user_id).single().execute()
    if not user_data.data or not user_data.data.get("cv_url"):
        return

    cv_path = user_data.data["cv_url"]

    try:
        signed = supabase.storage.from_("user-uploads").create_signed_url(cv_path, 3600)
        cv_signed_url = signed.get("signedURL") or signed.get("signedUrl")


        if not cv_signed_url:
            return

        file_response = requests.get(cv_signed_url)
        file_response.raise_for_status()
    except Exception:
        return

    try:
        ext = cv_path.split("?")[0].split(".")[-1].lower()
        cv_text = ""

        if ext == "pdf":
            with pdfplumber.open(io.BytesIO(file_response.content)) as pdf:
                for page in pdf.pages:
                    text = page.extract_text()
                    if text:
                        cv_text += text + "\n"

        elif ext == "docx":
            doc = Document(io.BytesIO(file_response.content))
            for paragraph in doc.paragraphs:
                cv_text += paragraph.text + "\n"
        else:
            return

    except Exception:
        return

    if not cv_text.strip():
        return

    job_description = job_data.data["description"]
    prompt = build_prompt(job_description, cv_text)

    try:
        model = genai.GenerativeModel("models/gemini-2.5-flash")
        response = model.generate_content(prompt)
        analysis_result = response.text if hasattr(response, "text") else str(response)
    except Exception as e:
        analysis_result = f"❌ Greška pri AI analizi: {str(e)}"

    score_match = re.search(r"(\d{1,2}\.\d)", analysis_result)
    score = float(score_match.group(1)) if score_match else 0.0

    supabase.table("application_analysis").update({
        "analysis": analysis_result,
        "score": score
    }).eq("id", analysis_id).execute()



def parse_pdf(file_path: str):
    text = ""
    with pdfplumber.open(file_path) as pdf:
        for page in pdf.pages:
            extracted = page.extract_text()
            if extracted:
                text += extracted + "\n"
    return text


def parse_image(image_path: str):
    image = Image.open(image_path)
    return pytesseract.image_to_string(image)


def parse_docx(file_path: str) -> str:
    text = ""
    try:
        doc = Document(file_path)
        for paragraph in doc.paragraphs:
            text += paragraph.text + "\n"
    except Exception as e:
        print(f"Greška pri čitanju .docx fajla: {e}")
    return text


@router.get("/get-analysis/{analysis_id}")
def get_analysis(analysis_id: str):
    try:
        result = supabase.table("application_analysis").select("analysis, score").eq("id", analysis_id).execute()
        if not result.data:
            raise HTTPException(status_code=404, detail="Nema analize")
        record = result.data[0]
        return {
            "analysis": record["analysis"],
            "score": record["score"],
        }
    except Exception:
        raise HTTPException(status_code=500, detail="Greška prilikom dohvata analize")


@router.get("/get-existing-analysis/{user_id}/{job_id}")
def get_existing_analysis(user_id: str, job_id: str):
    try:
        response = (
            supabase.table("application_analysis")
            .select("analysis")
            .eq("user_id", user_id)
            .eq("job_id", job_id)
            .execute()
        )

        if not response.data or len(response.data) == 0:
            return {"analysis": None}

        record = response.data[0]
        return {
            "analysis": record.get("analysis"),
        }
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Greška prilikom dohvata analize: {str(e)}")


def build_prompt(job_description: str, cv_text: str) -> str:
    return (
        f"📄 Opis posla:\n{job_description}\n\n"
        f"📄 CV kandidata:\n{cv_text}\n\n"
        "🔍 Sada slijedi analiza životopisa kandidata u odnosu na opis posla.\n\n"
        "📌 Upute za analizu:\n"
        "🔹 Na prvoj liniji **OBAVEZNO** napiši samo brojčanu ocjenu prikladnosti kandidata (npr. `7.5`).\n"
        "🔹 Koristi decimalnu vrijednost sa jednom decimalom (obavezno).\n"
        "🔹 Odgovor mora biti **100% zasnovan isključivo** na podacima iz CV-a i opisa posla. Nemoj izmišljati informacije koje nisu navedene.\n\n"
        "🔹 Nakon toga napiši **DETALJNU i STRUKTURIRANU** analizu kandidata.\n"
        "🔹 Tvoj odgovor treba biti organizovan u sljedeće sekcije (podnaslove):\n"
        "1. Kompetencije\n"
        "2. Iskustvo\n"
        "3. Edukacija\n"
        "4. Kompatibilnost\n"
        "5. Prednosti kandidata (navedi najmanje 3)\n"
        "6. Nedostaci kandidata (navedi najmanje 3)\n"
        "7. Preporuke\n\n"
        "🔹 Piši isključivo na **bosanskom jeziku**.\n"
        "🔹 Format odgovora:\n"
        "```\n"
        "8.3\n"
        "1. Kompetencije:\n...\n"
        "2. Iskustvo:\n...\n"
        "...\n"
        "```"
    )

# testing find my job 
@router.get("/find-my-jobs/{user_id}")
def find_my_jobs(user_id: str):
    try:
        result = supabase.table("users").select("cv_url").eq("id", user_id).single().execute()
        if not result.data or not result.data.get("cv_url"):
            raise HTTPException(status_code=404, detail="CV nije pronađen u bazi")

        cv_path = result.data["cv_url"]

        try:
            signed = supabase.storage.from_("user-uploads").create_signed_url(cv_path, 3600)
            cv_signed_url = signed.get("signedURL") or signed.get("signedUrl")


            if not cv_signed_url:
                raise HTTPException(status_code=500, detail="Ne mogu generisati URL za CV")

            file_response = requests.get(cv_signed_url)
            file_response.raise_for_status()
        except Exception as e:
            raise HTTPException(status_code=500, detail=f"Greška pri preuzimanju CV-a: {str(e)}")

        try:
            ext = cv_path.split("?")[0].split(".")[-1].lower()
            cv_text = ""

            if ext == "pdf":
                with pdfplumber.open(io.BytesIO(file_response.content)) as pdf:
                    for page in pdf.pages:
                        text = page.extract_text()
                        if text:
                            cv_text += text + "\n"

            elif ext == "docx":
                doc = Document(io.BytesIO(file_response.content))
                for paragraph in doc.paragraphs:
                    cv_text += paragraph.text + "\n"
            else:
                raise Exception("Nepodržan format fajla: ." + ext)

        except Exception as e:
            raise HTTPException(status_code=500, detail=f"Greška pri čitanju CV-a: {str(e)}")

        if not cv_text.strip():
            raise HTTPException(status_code=400, detail="CV je prazan")

        category = detect_cv_category(cv_text)

        prompt = f"""
        Tvoj zadatak je da automatski izvučeš najrelevantnije ključne riječi iz sljedećeg CV-a kandidata.

        Tekst CV-a:
        \"\"\"{cv_text}\"\"\"

        Pravila za izdvajanje ključnih riječi:
        1. Ključne riječi moraju biti pojedinačne riječi ili kratke fraze (1-3 riječi).
        2. Ne koristi rečenice, opise niti nepotrebne informacije.
        3. Vrati isključivo listu ključnih riječi u JSON formatu: ["riječ1", "riječ2", "fraza 3", ...]
        4. Ključne riječi treba da predstavljaju vještine, tehnologije, alate, pozicije ili industrijske izraze.
        5. Piši ih na jeziku na kojem su i u originalnom tekstu (npr. "React", "računovodstvo").

        Output mora izgledati ovako:
        ["javascript", "sql", "računovodstvo", "administracija", "komunikacija"]

        Nemoj dodavati nikakva pojašnjenja, samo vrati listu kao JSON niz.
        """

        model = genai.GenerativeModel("models/gemini-2.5-flash")
        response = model.generate_content(prompt)
        keywords_raw = response.text.strip()

        try:
            keywords = json.loads(keywords_raw)
        except:
            keywords = [kw.strip().lower() for kw in re.findall(r'\w+', keywords_raw)]

        if not keywords:
            raise HTTPException(status_code=400, detail="Nema pronađenih ključnih riječi")

        jobs_response = supabase.table("jobs").select("*").ilike("job_type", category).execute()
        jobs = jobs_response.data if jobs_response.data else []

        def match_score(job):
            combined_text = f"{job['title']} {job['description']}".lower()
            return sum(1 for kw in keywords if kw.lower() in combined_text)

        ranked_jobs = sorted(jobs, key=match_score, reverse=True)

        print("DETECTED CATEGORY:", category, flush=True)
        print("KEYWORDS:", keywords, flush=True)
        print("JOBS FOUND:", len(jobs), flush=True)

        return {
            "keywords": keywords,
            "category": category,
            "results": ranked_jobs[:10]
        }
    except Exception as e:
        print(" ERROR:", str(e), flush=True)
        raise HTTPException(status_code=500, detail=f"Greška: {str(e)}")


def detect_cv_category(cv_text: str) -> str:
    categories = ["it", "administracija", "ugostiteljstvo", "proizvodnja", "obrazovanje", "zdravstvo"]

    prompt = f"""
    Na osnovu sadržaja životopisa kandidata, odaberi samo jednu kategoriju zanimanja koja najbolje opisuje ovaj CV.

    Dostupne kategorije su:
    {', '.join(categories)}

    Tekst CV-a:
    \"\"\"
    {cv_text}
    \"\"\"

    Pravila:
    - Odgovori isključivo jednom od ponuđenih kategorija (ništa više).
    - Odgovor treba biti samo jedna riječ: it, administracija, ugostiteljstvo, proizvodnja, obrazovanje ili zdravstvo.
    - Ne dodaji objašnjenja.

    Samo rezultat:
    """

    try:
        model = genai.GenerativeModel("models/gemini-2.5-flash")
        response = model.generate_content(prompt)
        category = response.text.strip().lower()
        return category if category in categories else "nepoznato"
    except Exception as e:
        print(f"[Kategorija AI fallback] Greška: {e}")
        return "nepoznato"
